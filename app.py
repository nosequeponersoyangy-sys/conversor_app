# app.py

import streamlit as st
import io
import io
import os
from docx import Document
from docx.shared import Pt # Para usar en la función de PDF a Word
import fitz # Solo para verificar la disponibilidad

# Importar lógica modular
from ocr_utils import (
    configurar_tesseract, 
    extraer_texto_pdf_ocr, 
    limpiar_texto, 
    verificar_librerias,
    extraer_texto_documento
)
from tts_utils import generar_audio, VOCES

# Variables de estado de sesión
if 'texto_extraido' not in st.session_state:
    st.session_state['texto_extraido'] = ""
if 'nombre_archivo' not in st.session_state:
    st.session_state['nombre_archivo'] = ""

# Obtener disponibilidad de librerías al inicio
DOCX_OK, PYDUB_OK = verificar_librerias()

# ================= FUNCIÓN 1: PDF A WORD (OCR) =================
def vista_pdf_a_word(tesseract_ok):
    """Define la vista de PDF a Word."""
    st.subheader("📄 PDF → WORD (OCR)")
    st.markdown("Extrae texto de PDFs escaneados y guárdalo como documento Word.")
    
    if not DOCX_OK:
        st.error("La función PDF a WORD requiere 'python-docx'. Instálala: pip install python-docx")
        return
    if not tesseract_ok:
        st.warning("⚠️ Necesitas **Tesseract OCR** para esta función.")
        
    archivo_subido = st.file_uploader("Sube tu PDF", type="pdf", key="pdf_word")
    
    if archivo_subido:
        with st.form("ocr_form"):
            col1, col2 = st.columns(2)
            with col1:
                auto_rotar = st.checkbox("🔄 Enderezar páginas", True, help="Corrige la orientación de las páginas.")
            with col2:
                es_libro = st.checkbox("📖 Separar doble página", True, help="Divide páginas con formato de libro.")
            
            submit_button = st.form_submit_button("📝 Extraer texto a Word", type="primary", use_container_width=True)

        if submit_button:
            with st.spinner("Leyendo PDF con OCR..."):
                texto = extraer_texto_pdf_ocr(archivo_subido, es_libro, auto_rotar)
                
                if texto and len(texto.strip()) > 50:
                    # Crear documento Word
                    doc = Document()
                    doc.add_heading(f'Texto extraído de: {archivo_subido.name}', 0)
                    
                    # Agregar texto
                    parrafos = texto.split('\n\n')
                    for parrafo in parrafos:
                        if parrafo.strip():
                            p = doc.add_paragraph(parrafo.strip())
                            p.style.font.size = Pt(11)
                    
                    # Guardar temporalmente y descargar
                    nombre_base = os.path.splitext(archivo_subido.name)[0]
                    ruta_word = f"{nombre_base}_extraido.docx"
                    doc.save(ruta_word)
                    
                    with open(ruta_word, "rb") as f:
                        st.success("✅ Documento Word creado")
                        st.download_button(
                            "⬇️ Descargar Word (.docx)",
                            f.read(),
                            f"{nombre_base}_extraido.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    os.remove(ruta_word) # Limpieza
                    
                    # Mostrar previsualización y continuar a audio
                    with st.expander("📋 Ver texto extraído"):
                        st.text_area("Texto:", texto[:2000] + ("..." if len(texto) > 2000 else ""), height=200, key="prev_word_final")
                    
                    st.session_state['texto_extraido'] = texto
                    st.session_state['nombre_archivo'] = archivo_subido.name
                    if st.button("🎧 Continuar a Conversión de Audio", use_container_width=True, key="cont_audio"):
                        st.session_state['funcion'] = "🎧 PDF → AUDIO (Directo)"
                        st.rerun()
                else:
                    st.error("No se pudo extraer texto suficiente.")

# ================= FUNCIÓN 2: PDF A AUDIO (Directo) =================
def vista_pdf_a_audio(tesseract_ok):
    """Define la vista de PDF a Audio."""
    st.subheader("🎧 PDF → AUDIO (Directo)")
    st.markdown("Extrae texto con OCR y convierte directamente a audiolibro.")
    
    if not tesseract_ok:
        st.warning("⚠️ Necesitas **Tesseract OCR** para esta función.")
    if not PYDUB_OK:
        st.info("ℹ️ **pydub** no está instalado. El audio largo será truncado. Instala FFmpeg y pydub.")
        
    texto_a_usar = st.session_state['texto_extraido']
    nombre_base = st.session_state['nombre_archivo']
    
    if texto_a_usar:
        st.info(f"📝 Usando texto extraído previamente de: **{nombre_base}**")
        if st.button("🔄 Usar nuevo PDF en lugar del texto cargado", key="new_pdf"):
            st.session_state['texto_extraido'] = ""
            st.session_state['nombre_archivo'] = ""
            st.rerun()
    else:
        # Petición de archivo y extracción
        archivo_subido = st.file_uploader("Sube tu PDF", type="pdf", key="pdf_audio")
        if not archivo_subido:
            return
            
        with st.form("ocr_audio_form"):
            col1, col2 = st.columns(2)
            with col1:
                auto_rotar = st.checkbox("🔄 Enderezar páginas", True, key="auto2")
            with col2:
                es_libro = st.checkbox("📖 Separar doble página", True, key="libro2")
            
            submit_button = st.form_submit_button("📝 Extraer texto", type="secondary", use_container_width=True)
            
            if submit_button:
                with st.spinner("Procesando PDF..."):
                    texto_a_usar = extraer_texto_pdf_ocr(archivo_subido, es_libro, auto_rotar)
                    nombre_base = archivo_subido.name
                    st.session_state['texto_extraido'] = texto_a_usar
                    st.session_state['nombre_archivo'] = nombre_base
                    # st.rerun() # No es necesario si el resto del código se ejecuta después
                    
    if texto_a_usar and len(texto_a_usar.strip()) > 50:
        # --- Configuración de Audio ---
        st.markdown("---")
        st.subheader("🎵 Configurar Audio")
        
        col1, col2 = st.columns(2)
        with col1:
            voz_seleccionada = st.selectbox("Voz del narrador:", list(VOCES.keys()), key="voz2")
            voz_codigo = VOCES[voz_seleccionada]
        with col2:
            texto_limpio_temp = limpiar_texto(texto_a_usar)
            st.metric("📊 Caracteres", f"{len(texto_limpio_temp):,}")
            
        if st.button("🔊 Generar Audiolibro", type="primary", use_container_width=True, key="gen_audio"):
            texto_limpio_final = generar_audio(texto_a_usar, voz_codigo, nombre_base, PYDUB_OK)
            
            # Mostrar métricas después de la generación
            col_stats1, col_stats2 = st.columns(2)
            with col_stats1:
                st.metric("📝 Palabras", f"{len(texto_limpio_final.split()):,}")
            with col_stats2:
                tiempo_estimado = len(texto_limpio_final.split()) * 0.4 / 60
                st.metric("⏱️ Duración estimada", f"{tiempo_estimado:.1f} min")
                
            with st.expander("📋 Ver texto limpio"):
                st.text_area("Texto a convertir:", texto_limpio_final[:2000] + ("..." if len(texto_limpio_final) > 2000 else ""), height=200, key="prev_text_final")

# ================= FUNCIÓN 3: WORD/TEXTO A AUDIO =================
def vista_texto_a_audio():
    """Define la vista de Word/Texto a Audio."""
    st.subheader("📝 WORD/TEXTO → AUDIO")
    st.markdown("Convierte documentos (Word/TXT/PDF digital) o texto pegado a audiolibro.")

    texto_final = ""
    nombre_base = "texto_manual"

    # 1. Subir archivo
    archivo_subido = st.file_uploader(
        "Sube tu documento (.docx, .txt, .pdf)", 
        type=["docx", "txt", "pdf"], 
        key="texto_audio"
    )
    
    if archivo_subido:
        texto_final, nombre_original = extraer_texto_documento(archivo_subido)
        nombre_base = nombre_original
        if texto_final:
            st.success(f"✅ Archivo '{nombre_original}' cargado: {len(texto_final)} caracteres.")
        
    st.markdown("---")
    
    # 2. Pegar texto manualmente
    texto_manual = st.text_area("O pega tu texto aquí:", height=150, key="texto_manual")
    
    if texto_manual.strip() and not archivo_subido:
        texto_final = texto_manual
        st.success(f"✅ Texto manual cargado: {len(texto_final)} caracteres.")
    
    if texto_final and len(texto_final.strip()) > 50:
        texto_limpio = limpiar_texto(texto_final)
        
        # --- Configuración de Audio ---
        st.markdown("---")
        st.subheader("🎵 Configurar Audio")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            voz_seleccionada = st.selectbox("Voz del narrador:", list(VOCES.keys()), key="voz3")
            voz_codigo = VOCES[voz_seleccionada]
        with col2:
            st.metric("📊 Caracteres", f"{len(texto_limpio):,}")
        with col3:
            tiempo_estimado = len(texto_limpio.split()) * 0.4 / 60
            st.metric("⏱️ Audio estimado", f"{tiempo_estimado:.1f} min")
            
        # Previsualización
        with st.expander("📋 Ver texto limpio"):
            st.text_area("Texto a convertir:", texto_limpio[:2000] + ("..." if len(texto_limpio) > 2000 else ""), height=200, key="prev_text")

        if st.button("🔊 Convertir Texto a Audio", type="primary", use_container_width=True, key="conv_text_audio"):
            generar_audio(texto_final, voz_codigo, nombre_base, PYDUB_OK)

# ================= INTERFAZ PRINCIPAL =================
def main():
    st.set_page_config(
        page_title="Conversor Multifuncional OCR/Audio",
        page_icon="🎯",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Título y descripción
    st.title("🎯 CONVERSOR MULTIFUNCIONAL OCR/AUDIO")
    st.markdown("""
    **3 funciones en 1 aplicación:**
    1. 📄 **PDF → WORD**: Extrae texto de PDFs escaneados con OCR y guarda como Word
    2. 🎧 **PDF → AUDIO**: Convierte PDFs directamente a audiolibro (OCR + TTS)
    3. 📝 **WORD/TEXTO → AUDIO**: Convierte documentos existentes a audio
    """)
    
    # --- Configuración en Sidebar ---
    with st.sidebar:
        st.header("⚙️ Configuración y Estado")
        
        # 1. Tesseract
        tesseract_ok, tesseract_info = configurar_tesseract()
        if tesseract_ok:
            st.success(f"✅ Tesseract: {tesseract_info}")
        else:
            st.error(f"❌ Tesseract: No encontrado. Instálalo e inclúyelo en PATH.")
            st.info("💡 Necesario para funciones 1 y 2 (OCR).")
            
        # 2. python-docx
        if DOCX_OK:
            st.success("✅ python-docx: OK")
        else:
            st.warning("⚠️ python-docx: Faltante. Word a/desde PDF no disponible.")
        
        # 3. pydub
        if PYDUB_OK:
            st.success("✅ pydub (Audio largo): OK")
        else:
            st.warning("⚠️ pydub (Audio largo): Faltante. Audio largo será truncado.")
            st.info("💡 Necesitas **FFmpeg** instalado en tu sistema además de `pip install pydub`.")
        
        st.divider()
        st.subheader("📁 Funciones")
        
        # Selector de función (usa session_state si viene de "Continuar")
        funcion = st.radio(
            "Selecciona la función:",
            ["📄 PDF → WORD (OCR)", "🎧 PDF → AUDIO (Directo)", "📝 WORD/TEXTO → AUDIO"],
            index=0 if st.session_state['funcion'] == "" else ["📄 PDF → WORD (OCR)", "🎧 PDF → AUDIO (Directo)", "📝 WORD/TEXTO → AUDIO"].index(st.session_state['funcion']),
            key="funcion_radio"
        )
        st.session_state['funcion'] = funcion
        
        st.divider()
        st.info("🔄 **Flujo recomendado:** PDF (Escaneado) → WORD (OCR) → Audio")
    
    # --- Contenido principal ---
    st.markdown("---")
    
    if st.session_state['funcion'] == "📄 PDF → WORD (OCR)":
        vista_pdf_a_word(tesseract_ok)
    
    elif st.session_state['funcion'] == "🎧 PDF → AUDIO (Directo)":
        vista_pdf_a_audio(tesseract_ok)
    
    elif st.session_state['funcion'] == "📝 WORD/TEXTO → AUDIO":
        vista_texto_a_audio()

    # Pie de página
    st.markdown("---")
    st.caption("Hecho con Streamlit, PyMuPDF, Pytesseract, edge-tts y pydub.")

# ================= EJECUCIÓN =================
if __name__ == "__main__":
    if 'funcion' not in st.session_state:
        st.session_state['funcion'] = "📄 PDF → WORD (OCR)" # Valor inicial
    
    main()